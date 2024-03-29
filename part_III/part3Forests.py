import pandas as pd
import docx
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import confusion_matrix

f1_list = []  # list for f1-scores
p_list = []  # list for precision
r_list = []  # list for recall

tp_list = []
fp_list = []
fn_list = []
tn_list = []


def learn_dt(file_train, file_test, criterion, n_estimators, max_features, max_depth, random_state):
    # load the training data as a matrix
    dataset = pd.read_csv(file_train, header=0)

    # separate the data from the target attributes
    train_data = dataset.drop('500_Buggy?', axis=1)

    # remove unnecessary features
    train_data = train_data.drop('change_id', axis=1)
    train_data = train_data.drop('412_full_path', axis=1)
    train_data = train_data.drop('411_commit_time', axis=1)

    # the lables of training data. `label` is the title of the  last column in your CSV files
    train_target = dataset.loc[:, '500_Buggy?']

    # load the testing data
    dataset2 = pd.read_csv(file_test, header=0)

    # separate the data from the target attributes
    test_data = dataset2.drop('500_Buggy?', axis=1)

    # remove unnecessary features
    test_data = test_data.drop('change_id', axis=1)
    test_data = test_data.drop('412_full_path', axis=1)
    test_data = test_data.drop('411_commit_time', axis=1)

    # the lables of test data
    test_target = dataset2.loc[:, '500_Buggy?']

    r_forests = RandomForestClassifier(criterion=criterion, n_estimators=n_estimators, max_features=max_features,
                                       max_depth=max_depth, random_state=random_state)

    test_pred = r_forests.fit(train_data, train_target).predict(test_data)
    conf_matrix = confusion_matrix(test_target, test_pred, labels=[0, 1])
    TP = conf_matrix[0][0]
    FP = conf_matrix[0][1]
    FN = conf_matrix[1][0]
    TN = conf_matrix[1][1]
    tp_list.append(TP)
    fp_list.append(FP)
    fn_list.append(FN)
    tn_list.append(TN)


def param_test_jr(criterion, n_estimators, max_feat, max_d):
    for i in range(0, 6):
        learn_dt("../data/jackrabbit/" + str(i) + "/train_bow.csv", "../data/jackrabbit/" + str(i) + "/test_bow.csv",
                 criterion=criterion, n_estimators=n_estimators, max_features=max_feat,
                 max_depth=max_d, random_state=52)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


def param_test_jdt(criterion, n_estimators, max_feat, max_d):
    # jdt
    for i in range(0, 6):
        # print("Conducting tests on set " + str(i))
        learn_dt("../data/jdt/" + str(i) + "/train_bow.csv", "../data/jdt/" + str(i) + "/test_bow.csv",
                 criterion=criterion, n_estimators=n_estimators, max_features=max_feat,
                 max_depth=max_d, random_state=52)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


def param_test_lucene(criterion, n_estimators, max_feat, max_d):
    # lucene
    for i in range(0, 6):
        learn_dt("../data/lucene/" + str(i) + "/train_bow.csv", "../data/lucene/" + str(i) + "/test_bow.csv",
                 criterion=criterion, n_estimators=n_estimators, max_features=max_feat,
                 max_depth=max_d, random_state=52)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


def param_test_xorg(criterion, n_estimators, max_feat, max_d):
    # xorg
    for i in range(0, 6):
        # print("Conducting tests on set " + str(i))
        learn_dt("../data/xorg/" + str(i) + "/train_bow.csv", "../data/xorg/" + str(i) + "/test_bow.csv",
                 criterion=criterion, n_estimators=n_estimators, max_features=max_feat,
                 max_depth=max_d, random_state=52)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


doc = docx.Document("test.docx")


def to_doc(d_frame):
    # global doc
    t = doc.add_table(d_frame.shape[0] + 1, d_frame.shape[1])

    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])
    doc.save("test.docx")


n_est = [1, 2, 5, 8, 10, 20, 30, 50, 80, 100]
n_depth = [1, 2, 5, 8, 10, 20, 30, 50, 80, 100]
n_feat = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13]
#
doc.add_paragraph("Using Gini")
print("Using Gini...")
print("jackrabbit\n")
for x in n_est:
    param_test_jr("gini", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("jdt\n")
for x in n_est:
    param_test_jdt("gini", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("lucene\n")
for x in n_est:
    param_test_lucene("gini", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("xorg\n")
for x in n_est:
    param_test_xorg("gini", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

doc.add_paragraph("Using Entropy")
print("Using Entropy...\n")
print("jackrabbit\n")
for x in n_est:
    param_test_jr("entropy", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("jdt\n")
for x in n_est:
    param_test_jdt("entropy", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("lucene\n")
for x in n_est:
    param_test_lucene("entropy", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("xorg\n")
for x in n_est:
    param_test_xorg("entropy", x, None, None)
df = pd.DataFrame(list(zip(n_est, p_list, r_list, f1_list)),
                  columns=["n_estimators", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

doc.add_paragraph("Using Criterion = 'entropy' and n_estimators = 80 to find max depth...")
print("Using Criterion = 'entropy' and n_estimators = 80 to find max depth...\n")
print("jackrabbit\n")
for x in n_depth:
    param_test_jr("entropy", 80, None, x)
df = pd.DataFrame(list(zip(n_depth, p_list, r_list, f1_list)),
                  columns=["max_depth", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("jdt\n")
for x in n_depth:
    param_test_jdt("entropy", 80, None, x)
df = pd.DataFrame(list(zip(n_depth, p_list, r_list, f1_list)),
                  columns=["max_depth", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("lucene\n")
for x in n_depth:
    param_test_lucene("entropy", 80, None, x)
df = pd.DataFrame(list(zip(n_depth, p_list, r_list, f1_list)),
                  columns=["max_depth", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("xorg\n")
for x in n_depth:
    param_test_xorg("entropy", 80, None, x)
df = pd.DataFrame(list(zip(n_depth, p_list, r_list, f1_list)),
                  columns=["max_depth", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()
#
print("Using Criterion = 'entropy', n_estimators = 80 and max_depth = 30 to find max features...\n")
doc.add_paragraph("Using Criterion = 'gini', n_estimators = 20 and max_depth = 10 to find max features...")
print("jackrabbit\n")
for x in n_feat:
    param_test_jr("entropy", 80, x, 30)
df = pd.DataFrame(list(zip(n_feat, p_list, r_list, f1_list)),
                  columns=["max_feat", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("jdt\n")
for x in n_feat:
    param_test_jdt("entropy", 80, x, 30)
df = pd.DataFrame(list(zip(n_feat, p_list, r_list, f1_list)),
                  columns=["max_feat", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("lucene\n")
for x in n_feat:
    param_test_lucene("entropy", 80, x, 30)
df = pd.DataFrame(list(zip(n_feat, p_list, r_list, f1_list)),
                  columns=["max_feat", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("xorg\n")
for x in n_feat:
    param_test_xorg("entropy", 80, x, 30)
df = pd.DataFrame(list(zip(n_feat, p_list, r_list, f1_list)),
                  columns=["max_feat", "Precision", "Recall", "F1-Score"])
to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()
